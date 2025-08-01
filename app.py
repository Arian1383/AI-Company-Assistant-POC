import streamlit as st
import os
import json
import time
from datetime import datetime
import base64
import tempfile
import docx # For Word files
import pandas as pd # For Excel files
import requests # For making HTTP requests to custom AI API
import shutil # For removing directories

# LangChain and AI related imports
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.embeddings.base import Embeddings # Base class for custom embeddings
from langchain.llms.base import LLM # Base class for custom LLM
from typing import Any, List, Mapping, Optional # For type hinting in custom LLM/Embeddings
from pydantic import Field # NEW: Import Field for Pydantic validation

from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document # Import Document here for load_knowledge_base

# --- Constants ---
USERS_FILE = "users.json"
FAISS_INDEX_PATH = "faiss_index"
KNOWLEDGE_SOURCES_DIR = "knowledge_source_files" # Directory to store all source files
CSS_FILE_LIGHT = "style_light.css" # نام فایل CSS برای تم روشن
CSS_FILE_DARK = "style_dark.css" # نام فایل CSS برای تم تیره

# --- Page Configuration (MUST be the first Streamlit command) ---
st.set_page_config(
    page_title="دستیار دانش شرکت سپاهان",
    page_icon="🤖",
    layout="wide", # استفاده از layout="wide" برای فضای بیشتر
    initial_sidebar_state="collapsed"
)

# --- Session State Initialization ---
def initialize_session_state():
    """Initializes session state variables if they don't exist."""
    defaults = {
        "initialized": True,
        "current_page": "login",
        "authenticated": False,
        "is_admin": False,
        "user_id": None,
        "theme": "light", # تم پیش‌فرض
        "messages": [{"role": "assistant", "type": "text", "content": "سلام! من دستیار هوشمند گروه صنعتی سپاهان هستم. در این سامانه می‌توانید سوالات خود را در مورد دستورالعمل‌ها و رویه‌های شرکت بپرسید و پاسخ فوری دریافت کنید."}],
        "page_history": [] # New: To store navigation history
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

initialize_session_state()

# --- CSS and Theme Management ---
def set_theme(theme_name):
    """Sets the theme in session state. Streamlit will rerun automatically."""
    st.session_state.theme = theme_name
    # Removed st.rerun() here as it's a no-op in on_change callbacks

def load_and_inject_css():
    """Reads the current theme's CSS file and injects it into the app."""
    css_file_path = CSS_FILE_LIGHT if st.session_state.theme == "light" else CSS_FILE_DARK
    
    if os.path.exists(css_file_path):
        with open(css_file_path, "r", encoding="utf-8") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    else:
        st.warning(f"⚠️ فایل CSS '{css_file_path}' پیدا نشد. لطفاً آن را در کنار 'app.py' قرار دهید.")
    
    # Inject general Streamlit overrides and font
    st.markdown(f"""
        <style>
        /* پنهان کردن هدر پیش‌فرض Streamlit و فوتر "Made with Streamlit" */
        header {{ visibility: hidden; }}
        .stApp footer {{ visibility: hidden; }}
        /* اعمال فونت و جهت متن به بدنه اصلی Streamlit */
        body {{
            font-family: 'Vazirmatn', sans-serif;
            direction: rtl;
            text-align: right;
            transition: background-color 0.3s ease, color 0.3s ease;
        }}
        </style>
    """, unsafe_allow_html=True)

load_and_inject_css() # Call this after session state is initialized

# --- Global Theme Switcher (positioned using CSS) ---
is_dark_mode = st.session_state.theme == "dark"
st.toggle("🌙", value=is_dark_mode, key="global_theme_toggle", help="تغییر تم (روشن/تیره)", label_visibility="hidden", on_change=lambda: set_theme("dark" if not is_dark_mode else "light"))

st.markdown(f"""
    <style>
    div[data-testid="stToggle"] {{
        position: fixed;
        top: 10px;
        left: 10px;
        z-index: 9999;
        margin: 0 !important;
        padding: 0 !important;
        background-color: var(--secondary-bg-color);
        border-radius: 15px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.2);
        display: flex;
        align-items: center;
        justify-content: center;
        width: 60px;
        height: 30px;
    }}
    div[data-testid="stToggle"] .st-bo {{
        width: 50px;
        height: 25px;
        border-radius: 15px;
        background-color: var(--input-bg);
        border: 1px solid var(--border-color);
        transition: all 0.3s ease;
        position: relative;
    }}
    div[data-testid="stToggle"] .st-bo > div:first-child {{
        background-color: transparent;
        border-radius: 15px;
    }}
    div[data-testid="stToggle"] .st-bo > div:first-child > div:first-child {{
        width: 20px;
        height: 20px;
        background-color: var(--accent-color);
        border-radius: 50%;
        margin: 2px;
        transition: all 0.3s ease;
        display: flex;
        align-items: center;
        justify-content: center;
        background-size: 80%;
        background-repeat: no-repeat;
        background-position: center;
    }}
    div[data-testid="stToggle"] input:not(:checked) + div > div > div {{
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='%23{f'{int(0.05 * 255):02x}'}{f'{int(0.05 * 255):02x}'}{f'{int(0.05 * 255):02x}'}'%3E%3Cpath d='M12 2.5a.5.5 0 01.5.5v2a.5.5 0 01-1 0v-2a.5.5 0 01.5-.5zM12 19.5a.5.5 0 01.5.5v2a.5.5 0 01-1 0v-2a.5.5 0 01.5-.5zM19.5 12a.5.5 0 01.5.5h2a.5.5 0 010 1h-2a.5.5 0 01-.5-.5v-1a.5.5 0 01.5-.5zM2.5 12a.5.5 0 01.5.5h2a.5.5 0 010 1h-2a.5.5 0 01-.5-.5v-1a.5.5 0 01.5-.5zM17.657 6.343a.5.5 0 01.354.146l1.414 1.414a.5.5 0 01-.707.707l-1.414-1.414a.5.5 0 01.353-.853zM4.929 17.657a.5.5 0 01.354.146l1.414 1.414a.5.5 0 01-.707.707l-1.414-1.414a.5.5 0 01.353-.853zM17.657 17.657a.5.5 0 01.354.146l1.414 1.414a.5.5 0 01-.707.707l-1.414-1.414a.5.5 0 01.353-.853zM4.929 6.343a.5.5 0 01.354.146l1.414 1.414a.5.5 0 01-.707.707l-1.414-1.414a.5.5 0 01.353-.853zM12 8a4 4 0 100 8 4 4 0 000-8z'/%3E%3C/svg%3E");
        background-color: var(--accent-color);
        transform: translateX(25px);
    }}
    div[data-testid="stToggle"] input:checked + div > div > div {{
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='%23{f'{int(0.9 * 255):02x}'}{f'{int(0.9 * 255):02x}'}{f'{int(0.9 * 255):02x}'}'%3E%3Cpath d='M12.3 4.5c.3-.3.7-.5 1.1-.5h.1c.4 0 .8.2 1.1.5.3.3.5.7.5 1.1v.1c0 .4-.2.8-.5 1.1-.3.3-.7.5-1.1.5h-.1c-.4 0-.8-.2-1.1-.5-.3-.3-.5-.7-.5-1.1v-.1c0-.4.2-.8.5-1.1zM12 2a10 10 0 100 20 10 10 0 000-20zM12 4a8 8 0 110 16 8 8 0 010-16zM13 5c-3.866 0-7 3.134-7 7s3.134 7 7 7c.302 0 .598-.02.89-.06a.5.5 0 01.61.61c-.3.292-.61.573-.93.837-1.17.96-2.61 1.52-4.17 1.52-4.97 0-9-4.03-9-9s4.03-9 9-9c1.56 0 3 .56 4.17 1.52.32.264.63.545.93.837a.5.5 0 01-.61.61c-.29-.04-.58-.06-.89-.06z'/%3E%3C/svg%3E");
        background-color: var(--accent-color);
        transform: translateX(0);
    }}
    div[data-testid="stToggle"] label > div:last-child {{
        display: none;
    }}
    </style>
""", unsafe_allow_html=True)


# --- API Key Management ---
try:
    # Use a generic name for the custom AI API key
    aval_ai_api_key = st.secrets["AVAL_AI_API_KEY"] 
except KeyError:
    st.error("🔑 خطای کلید API: کلید 'AVAL_AI_API_KEY' پیدا نشد. لطفاً آن را در Streamlit Secrets تنظیم کنید.")
    st.stop()


# --- Custom LLM and Embeddings Classes ---
# IMPORTANT: You need to implement the actual API calls for your chosen AI model here.
# The `api_url` and `model` names are placeholders and MUST be replaced with Aval AI's actual endpoints and model names.

class CustomEmbeddings(Embeddings):
    """Custom Embeddings model for Aval AI."""
    # Pydantic fields defined directly as class attributes
    api_key: str = Field(description="API Key for Aval AI Embeddings")
    # IMPORTANT: Replace with the actual API URL for Aval AI's embedding endpoint
    api_url: str = Field("https://api.aval.ai/v1/embeddings", description="API URL for Aval AI Embeddings endpoint")
    model_name: str = Field("text-embedding-ada-002", description="Model name for Aval AI Embeddings") # Example model name, replace with Aval's embedding model name

    # Explicit __init__ to ensure Pydantic fields are correctly passed to superclass
    # This __init__ is crucial for Pydantic v2+ when inheriting from BaseModel or similar
    def __init__(self, **data: Any):
        super().__init__(**data)


    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed a list of texts using Aval AI API."""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        # IMPORTANT: Adjust payload as per Aval AI's embedding API documentation
        payload = {
            "input": texts,
            "model": self.model_name # Use the model_name defined
        }
        
        try:
            response = requests.post(self.api_url, headers=headers, json=payload)
            response.raise_for_status() # Raise an exception for HTTP errors (4xx or 5xx)
            
            # IMPORTANT: Adjust parsing based on Aval AI's API response structure
            # Assuming response.json() looks like {"data": [{"embedding": [...]}, ...]}
            embeddings_data = response.json().get("data", [])
            if not embeddings_data:
                st.error("API Embeddings 'اول' پاسخ معتبری برنگرداند.")
                # Return a list of zero-filled vectors matching the number of texts
                # and a common embedding dimension (e.g., 768 or 1536 for common models)
                return [[0.0] * 768 for _ in texts] 
            
            # Extract embeddings from the data list
            extracted_embeddings = [item["embedding"] for item in embeddings_data if "embedding" in item]
            
            if not extracted_embeddings:
                st.error("Embeddings معتبری در پاسخ API 'اول' یافت نشد.")
                return [[0.0] * 768 for _ in texts]

            return extracted_embeddings
        except requests.exceptions.RequestException as e:
            st.error(f"خطا در فراخوانی API Embeddings 'اول': {e}")
            return [[0.0] * 768 for _ in texts] # Return dummy embeddings on error
        except Exception as e:
            st.error(f"خطا در پردازش پاسخ API Embeddings 'اول': {e}")
            return [[0.0] * 768 for _ in texts]

    def embed_query(self, text: str) -> List[float]:
        """Embed a single query text using Aval AI API."""
        return self.embed_documents([text])[0]

    @property
    def _identifying_params(self) -> Mapping[str, Any]:
        return {"api_key": self.api_key, "api_url": self.api_url, "model_name": self.model_name}

    @property
    def _llm_type(self) -> str:
        return "aval-embeddings"


class CustomLLM(LLM):
    """Custom LLM for Aval AI."""
    # Pydantic fields defined directly as class attributes
    api_key: str = Field(description="API Key for Aval AI LLM")
    # IMPORTANT: Replace with the actual API URL for Aval AI's chat endpoint
    api_url: str = Field("https://api.aval.ai/v1/chat/completions", description="API URL for Aval AI Chat endpoint")
    model_name: str = Field("gpt-3.5-turbo", description="Model name for Aval AI Chat") # Example model name, replace with Aval's chat model name

    # Explicit __init__ to ensure Pydantic fields are correctly passed to superclass
    def __init__(self, **data: Any):
        super().__init__(**data)

    @property
    def _llm_type(self) -> str:
        return "aval-ai"

    def _call(self, prompt: str, stop: Optional[List[str]] = None) -> str:
        """Call out to Aval AI's LLM API."""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        # IMPORTANT: Adjust payload as per Aval AI's chat API documentation
        # Assuming a chat-like API structure
        payload = {
            "model": self.model_name,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 500 # Example, adjust as needed
        }

        try:
            response = requests.post(self.api_url, headers=headers, json=payload)
            response.raise_for_status()
            
            # IMPORTANT: Adjust this line based on how Aval AI's API returns the generated text
            # Assuming response.json() looks like {"choices": [{"message": {"content": "..."}}]}
            generated_text = response.json().get("choices", [{}])[0].get("message", {}).get("content", "No response from AI.")
            return generated_text
        except requests.exceptions.RequestException as e:
            st.error(f"خطا در فراخوانی API LLM 'اول': {e}")
            return "متاسفانه در حال حاضر مشکلی در پردازش درخواست شما پیش آمده است."
        except Exception as e:
            st.error(f"خطا در پردازش پاسخ API LLM 'اول': {e}")
            return "متاسفانه در حال حاضر مشکلی در پردازش درخواست شما پیش آمده است."

    @property
    def _identifying_params(self) -> Mapping[str, Any]:
        return {"api_key": self.api_key, "api_url": self.api_url, "model_name": self.model_name}


# --- CORE LOGIC ---

def load_users():
    """Loads user data from JSON file. Creates default users if file doesn't exist."""
    if not os.path.exists(USERS_FILE):
        default_users = {
            "users": [{"username": "Sepahan", "password": "Arian"}],
            "admin_users": [{"username": "admin_sepahan", "password": "admin_pass"}]
        }
        with open(USERS_FILE, "w", encoding="utf-8") as f:
            json.dump(default_users, f, indent=4)
        return default_users
    with open(USERS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_users(users_data):
    """Saves user data to JSON file."""
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users_data, f, indent=4)

def navigate_to(page_name):
    """Navigates to a new page, pushing current page to history."""
    if st.session_state.current_page != page_name:
        st.session_state.page_history.append(st.session_state.current_page)
        st.session_state.current_page = page_name
        st.rerun() # Keep rerun here as it's a direct navigation action

def go_back():
    """Navigates back to the previous page in history."""
    if st.session_state.page_history:
        st.session_state.current_page = st.session_state.page_history.pop()
        st.rerun() # Keep rerun here as it's a direct navigation action
    else:
        st.warning("هیچ صفحه قبلی برای بازگشت وجود ندارد.")

def go_to_main_page():
    """Navigates to the user's primary page (chat for regular user, admin for admin)."""
    if st.session_state.is_admin:
        navigate_to("admin")
    else:
        navigate_to("chat")

def validate_credentials(username, password, is_admin_attempt=False):
    """Validates user credentials and updates session state."""
    users_data = load_users()
    user_type_key = "admin_users" if is_admin_attempt else "users"
    
    for user_info in users_data.get(user_type_key, []):
        if user_info["username"] == username and user_info["password"] == password:
            st.session_state.user_id = username
            st.session_state.authenticated = True
            st.session_state.is_admin = is_admin_attempt
            
            target_page = "admin" if is_admin_attempt else "chat"
            navigate_to(target_page) 
            st.success("✅ ورود موفقیت‌آمیز! در حال انتقال...")
            time.sleep(1)
            return True
    st.error("❌ نام کاربری یا رمز عبور اشتباه است.")
    return False

def logout():
    """Logs out the current user and resets session state."""
    theme_before_logout = st.session_state.theme # Preserve theme
    st.session_state.clear()
    initialize_session_state()
    st.session_state.theme = theme_before_logout # Restore theme
    st.rerun()

def create_user(username, password, is_admin):
    """Creates a new user (admin or regular)."""
    if not username or not password:
        st.warning("لطفاً نام کاربری و رمز عبور را وارد کنید.")
        return
    users_data = load_users()
    if any(u["username"] == username for u in users_data["users"]) or \
       any(u["username"] == username for u in users_data["admin_users"]):
        st.warning("⚠️ کاربری با این نام کاربری از قبل وجود دارد.")
        return
    
    target_list = "admin_users" if is_admin else "users"
    users_data[target_list].append({"username": username, "password": password})
    save_users(users_data)
    st.success(f"✅ کاربر '{username}' با موفقیت ایجاد شد.")
    time.sleep(1)
    st.rerun()

def delete_user(username_to_delete):
    """Deletes a user by username."""
    users_data = load_users()
    deleted = False
    for user_type in ["users", "admin_users"]:
        initial_len = len(users_data[user_type])
        users_data[user_type] = [u for u in users_data[user_type] if u['username'] != username_to_delete]
        if len(users_data[user_type]) < initial_len:
            save_users(users_data)
            deleted = True
            break
    if deleted:
        st.success(f"✅ کاربر '{username_to_delete}' با موفقیت حذف شد.")
        time.sleep(1)
        st.rerun()
    else:
        st.warning("⚠️ کاربری با این نام کاربری یافت نشد.")

@st.cache_resource(ttl=3600)
def load_knowledge_base_from_index(api_key_for_embeddings): # Renamed parameter for clarity
    """Loads the FAISS knowledge base from disk."""
    if not os.path.exists(FAISS_INDEX_PATH):
        st.warning("⚠️ پوشه پایگاه دانش FAISS یافت نشد. لطفاً ابتدا فایل‌های پایگاه دانش را بارگذاری کنید.")
        return None, None # Return None for both vector_store and chunks
    
    try:
        # Use CustomEmbeddings here
        embeddings = CustomEmbeddings(api_key=api_key_for_embeddings)
        
        vector_store = FAISS.load_local(FAISS_INDEX_PATH, embeddings, allow_dangerous_deserialization=True)
        return vector_store, None 
    except Exception as e:
        st.error(f"🚨 خطایی در بارگذاری پایگاه دانش رخ داد: {e}")
        return None, None

def process_file_to_documents(file_path, file_extension):
    """Processes a single file (PDF, DOCX, XLSX) and returns a list of LangChain Document objects."""
    documents = []
    if file_extension == ".pdf":
        loader = PyPDFLoader(file_path)
        documents.extend(loader.load())
    elif file_extension == ".docx":
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        documents.append(Document(page_content="\n".join(full_text)))
    elif file_extension == ".xlsx":
        xls = pd.ExcelFile(file_path)
        full_text = []
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)
            full_text.append(f"Sheet: {sheet_name}\n{df.to_string(index=False)}")
        documents.append(Document(page_content="\n".join(full_text)))
    return documents

def rebuild_knowledge_base(api_key_for_embeddings): # Pass API key for embeddings
    """Rebuilds the FAISS knowledge base from all files in KNOWLEDGE_SOURCES_DIR."""
    st.info(f"مسیر پوشه منابع دانش: {KNOWLEDGE_SOURCES_DIR}")
    # Ensure KNOWLEDGE_SOURCES_DIR exists
    if not os.path.exists(KNOWLEDGE_SOURCES_DIR):
        try:
            os.makedirs(KNOWLEDGE_SOURCES_DIR)
            st.info(f"پوشه منابع دانش '{KNOWLEDGE_SOURCES_DIR}' ایجاد شد.")
        except OSError as e:
            st.error(f"❌ خطای دسترسی: قادر به ایجاد پوشه منابع دانش '{KNOWLEDGE_SOURCES_DIR}' نیستم. لطفاً مجوزهای پوشه اصلی را بررسی کنید. خطا: {e}")
            return
    else:
        st.info(f"پوشه منابع دانش '{KNOWLEDGE_SOURCES_DIR}' از قبل وجود دارد.")

    # --- NEW: Delete existing FAISS index to ensure a fresh build ---
    if os.path.exists(FAISS_INDEX_PATH):
        try:
            st.info(f"در حال حذف پوشه پایگاه دانش FAISS قبلی: {FAISS_INDEX_PATH}")
            shutil.rmtree(FAISS_INDEX_PATH)
            st.info("پوشه پایگاه دانش FAISS قبلی با موفقیت حذف شد.")
        except Exception as e:
            st.warning(f"⚠️ خطایی در حذف پوشه FAISS رخ داد: {e}. ممکن است پایگاه دانش به درستی بازسازی نشود. لطفاً مجوزها را بررسی کنید.")
    else:
        st.info(f"پوشه پایگاه دانش FAISS در مسیر '{FAISS_INDEX_PATH}' یافت نشد. یک ایندکس جدید ایجاد خواهد شد.")
    # --- END NEW ---

    all_documents = []
    processed_files_count = 0
    
    # Iterate through all files in the source directory
    st.info(f"در حال اسکن فایل‌ها در پوشه: {KNOWLEDGE_SOURCES_DIR}")
    files_in_source_dir = []
    try:
        files_in_source_dir = [f for f in os.listdir(KNOWLEDGE_SOURCES_DIR) if os.path.isfile(os.path.join(KNOWLEDGE_SOURCES_DIR, f))]
    except OSError as e:
        st.error(f"❌ خطای دسترسی: قادر به خواندن پوشه منابع دانش '{KNOWLEDGE_SOURCES_DIR}' نیستم. لطفاً مجوزها را بررسی کنید. خطا: {e}")
        st.cache_resource.clear()
        return

    if not files_in_source_dir:
        st.warning("هیچ فایلی در پوشه منابع پایگاه دانش یافت نشد. پایگاه دانش بازسازی نشد. لطفاً فایل‌هایی را بارگذاری کنید.")
        st.cache_resource.clear()
        return

    for filename in files_in_source_dir:
        file_path = os.path.join(KNOWLEDGE_SOURCES_DIR, filename)
        file_extension = os.path.splitext(filename)[1].lower()

        if file_extension in [".pdf", ".docx", ".xlsx"]:
            try:
                st.info(f"در حال پردازش فایل: {filename}")
                documents = process_file_to_documents(file_path, file_extension)
                all_documents.extend(documents)
                processed_files_count += 1
                st.info(f"فایل '{filename}' با موفقیت پردازش شد.")
            except Exception as e:
                st.warning(f"⚠️ خطایی در پردازش فایل '{filename}' رخ داد: {e}. این فایل نادیده گرفته شد.")
        else:
            st.info(f"فایل '{filename}' نادیده گرفته شد. فرمت آن پشتیبانی نمی‌شود.")
    
    if not all_documents:
        st.warning("هیچ سند قابل پردازشی در پوشه منابع پایگاه دانش یافت نشد. پایگاه دانش بازسازی نشد.")
        st.cache_resource.clear()
        return

    st.info(f"تعداد کل اسناد پردازش شده: {len(all_documents)}")
    st.info("در حال تقسیم اسناد به بخش‌های کوچکتر (chunks)...")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200) 
    chunks = text_splitter.split_documents(all_documents)
    st.info(f"تعداد کل بخش‌های ایجاد شده: {len(chunks)}")
    
    st.info("در حال ایجاد Embeddings و ساخت پایگاه دانش FAISS...")
    try:
        embeddings = CustomEmbeddings(api_key=api_key_for_embeddings)
        vector_store = FAISS.from_documents(chunks, embeddings)
    except Exception as e:
        st.error(f"❌ خطایی در ایجاد Embeddings یا پایگاه دانش FAISS رخ داد: {e}")
        st.cache_resource.clear()
        return

    st.info(f"در حال ذخیره پایگاه دانش FAISS در مسیر: {FAISS_INDEX_PATH}")
    try:
        vector_store.save_local(FAISS_INDEX_PATH)
        # Add a small delay to ensure file system sync
        time.sleep(1) # Added sleep here
        if not os.path.exists(FAISS_INDEX_PATH):
            st.error("❌ هشدار: پوشه پایگاه دانش FAISS پس از ذخیره سازی پیدا نشد. ممکن است مشکل مجوز یا فضای دیسک وجود داشته باشد. لطفاً دسترسی‌های پوشه را بررسی کنید.")
        else:
            st.info(f"پایگاه دانش FAISS با موفقیت در '{FAISS_INDEX_PATH}' ذخیره شد.")
    except Exception as e:
        st.error(f"❌ خطایی در ذخیره سازی پایگاه دانش FAISS رخ داد: {e}. لطفاً مجوزهای نوشتن در پوشه را بررسی کنید.")
        st.cache_resource.clear()
        return
    
    st.cache_resource.clear() # Clear cache so load_knowledge_base_from_index reloads with new data
    return processed_files_count # Return count of processed files


def delete_knowledge_file(filename):
    """Deletes a file from KNOWLEDGE_SOURCES_DIR and rebuilds the knowledge base."""
    file_path = os.path.join(KNOWLEDGE_SOURCES_DIR, filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            st.success(f"✅ فایل '{filename}' با موفقیت حذف شد.")
            rebuild_knowledge_base(aval_ai_api_key) # Rebuild after deletion, pass API key
            st.rerun() # Rerun to update the file list and KB status
        except Exception as e:
            st.error(f"❌ خطایی در حذف فایل '{filename}' رخ داد: {e}. لطفاً مجوزهای دسترسی را بررسی کنید.")
    else:
        st.warning(f"⚠️ فایل '{filename}' یافت نشد.")


# --- UI RENDERING FUNCTIONS ---

def render_login_page():
    """Renders the login page."""
    _, center_col, _ = st.columns([1, 1.2, 1])
    with center_col:
        st.markdown('<h2 class="login-title">دستیار دانش گروه صنعتی سپاهان</h2>', unsafe_allow_html=True)
        st.markdown('<p class="login-subtitle">برای شروع، لطفاً با نام کاربری و رمز عبور خود وارد شوید. در صورت نداشتن حساب کاربری، با مدیر سیستم تماس بگیرید.</p>', unsafe_allow_html=True)
        
        login_tab, admin_tab = st.tabs(["ورود کاربر", "ورود مدیر"])
        with login_tab:
            with st.form("user_login_form"):
                username = st.text_input("نام کاربری", placeholder="نام کاربری خود را وارد کنید", label_visibility="collapsed")
                password = st.text_input("رمز عبور", type="password", placeholder="رمز عبور خود را وارد کنید", label_visibility="collapsed")
                if st.form_submit_button("ورود", use_container_width=True):
                    validate_credentials(username, password, is_admin_attempt=False)
        with admin_tab:
            with st.form("admin_login_form"):
                admin_username = st.text_input("نام کاربری مدیر", placeholder="نام کاربری ادمین", label_visibility="collapsed")
                admin_password = st.text_input("رمز عبور مدیر", type="password", placeholder="رمز عبور ادمین", label_visibility="collapsed")
                if st.form_submit_button("ورود مدیر", use_container_width=True):
                    validate_credentials(admin_username, admin_password, is_admin_attempt=True)
        st.markdown('</div>', unsafe_allow_html=True) # Closing login-card div


def render_admin_page():
    """Renders the admin panel page."""
    with st.sidebar:
        st.title(f"پنل مدیریت")
        st.caption(f"کاربر: {st.session_state.user_id}")
        
        st.markdown("---")
        if st.session_state.page_history:
            st.sidebar.button("🔙 بازگشت به صفحه قبلی", on_click=go_back, use_container_width=True)
        st.sidebar.button("🏠 بازگشت به صفحه اصلی", on_click=go_to_main_page, use_container_width=True)
        
        if st.sidebar.button("⚙️ حساب کاربری من", key="my_account_btn_admin", use_container_width=True):
            navigate_to("user_account")
        st.button("خروج از سیستم 🚪", on_click=logout, use_container_width=True)

    st.title("🛠️ مدیریت سیستم")
    
    admin_tabs = st.tabs(["📚 مدیریت پایگاه دانش", "👤 مدیریت کاربران", "📊 لاگ‌های سیستم"])

    with admin_tabs[0]:
        st.subheader("به‌روزرسانی پایگاه دانش")
        st.info("در این بخش می‌توانید فایل‌های PDF، Word و Excel را برای ساخت یا به‌روزرسانی پایگاه دانش بارگذاری کنید.")
        
        uploaded_files = st.file_uploader(
            "فایل‌های جدید را بارگذاری کنید (PDF, DOCX, XLSX)",
            type=["pdf", "docx", "xlsx"],
            accept_multiple_files=True, # Allow multiple files
            label_visibility="collapsed"
        )
        
        if uploaded_files:
            if st.button("🚀 افزودن و بازسازی پایگاه دانش", use_container_width=True, type="primary"):
                if not os.path.exists(KNOWLEDGE_SOURCES_DIR):
                    os.makedirs(KNOWLEDGE_SOURCES_DIR)
                
                files_saved_count = 0
                for uploaded_file in uploaded_files:
                    file_path = os.path.join(KNOWLEDGE_SOURCES_DIR, uploaded_file.name)
                    # Save the file
                    with open(file_path, "wb") as f:
                        f.write(uploaded_file.getvalue())
                    files_saved_count += 1
                
                st.info(f"✅ {files_saved_count} فایل جدید ذخیره شد. در حال بازسازی پایگاه دانش...")
                
                progress_bar = st.progress(0, text="در حال آماده‌سازی...")
                try:
                    processed_count = rebuild_knowledge_base(aval_ai_api_key) # Pass API key
                    progress_bar.progress(100, text="عملیات با موفقیت انجام شد!")
                    time.sleep(2)
                    st.success(f"✅ پایگاه دانش با موفقیت با {processed_count} سند به‌روزرسانی شد!")
                    st.balloons()
                    progress_bar.empty()
                except Exception as e:
                    progress_bar.empty()
                    st.error(f"❌ خطایی در هنگام بازسازی پایگاه دانش رخ داد: {e}")
                st.rerun() # Rerun to show updated file list

        st.markdown("<h5 class='admin-section-info'>فایل‌های موجود در پایگاه دانش:</h5>", unsafe_allow_html=True)
        if os.path.exists(KNOWLEDGE_SOURCES_DIR) and os.listdir(KNOWLEDGE_SOURCES_DIR):
            files_in_kb = [f for f in os.listdir(KNOWLEDGE_SOURCES_DIR) if os.path.isfile(os.path.join(KNOWLEDGE_SOURCES_DIR, f))]
            if files_in_kb:
                for file_name in files_in_kb:
                    file_path = os.path.join(KNOWLEDGE_SOURCES_DIR, file_name)
                    col1, col2, col3 = st.columns([0.6, 0.2, 0.2])
                    col1.write(f"**{file_name}**")
                    col2.write(f"{os.path.getsize(file_path) / (1024*1024):.2f} MB")
                    if col3.button("حذف", key=f"delete_file_{file_name}", use_container_width=True):
                        delete_knowledge_file(file_name) # This function handles rerun
            else:
                st.info("هیچ فایلی در پایگاه دانش یافت نشد. لطفاً فایل‌هایی را بارگذاری کنید.")
        else:
            st.info("هیچ فایلی در پایگاه دانش یافت نشد. لطفاً فایل‌هایی را بارگذاری کنید.")


    with admin_tabs[1]:
        st.subheader("مدیریت کاربران")
        st.info("در این بخش می‌توانید کاربران عادی و مدیر را اضافه یا حذف کنید.")
        with st.form("create_user_form"):
            cols = st.columns([2, 2, 1])
            new_user = cols[0].text_input("نام کاربری جدید", key="new_user_input")
            new_pass = cols[1].text_input("رمز عبور جدید", type="password", key="new_pass_input")
            is_admin_checkbox = cols[2].checkbox("مدیر باشد؟", key="is_admin_checkbox")
            if st.form_submit_button("ایجاد کاربر", use_container_width=True):
                create_user(new_user, new_pass, is_admin_checkbox)

        st.subheader("لیست کاربران موجود")
        users = load_users()
        all_users = users.get("users", []) + users.get("admin_users", [])
        if not all_users:
            st.info("هیچ کاربری یافت نشد.")
        else:
            for user in all_users:
                cols = st.columns([0.6, 0.2, 0.2])
                cols[0].write(f"**{user['username']}**")
                user_type = "مدیر" if user in users.get("admin_users", []) else "عادی"
                cols[1].markdown(f'<span class="user-role-badge role-{user_type.lower()}">{user_type}</span>', unsafe_allow_html=True)
                if user['username'] != st.session_state.user_id: # مدیر نتواند خودش را حذف کند
                    if cols[2].button("حذف", key=f"del_{user['username']}", use_container_width=True):
                        delete_user(user['username'])
                else:
                    cols[2].markdown("<span style='color: grey; font-size: 0.8em;'> (خودتان)</span>", unsafe_allow_html=True)


    with admin_tabs[2]:
        st.subheader("لاگ‌های مکالمات کاربران")
        st.info("در این بخش می‌توانید تاریخچه مکالمات کاربران با دستیار هوشمند را مشاهده کنید.")
        st.warning("⚠️ **توجه:** لاگ‌ها در نسخه POC بدون دیتابیس آنلاین، دائمی نیستند و با هر بار دیپلوی یا ری‌استارت برنامه از بین می‌روند.")
        
        st.markdown("""
        <div style="background-color: var(--secondary-bg-color); padding: 15px; border-radius: 10px; margin-bottom: 10px; color: var(--text-color); border: 1px solid var(--border-color);">
            <p style="font-size: 14px; margin-bottom: 5px;"><strong>لاگ‌ها در این نسخه ذخیره نمی‌شوند.</strong></p>
            <p style="font-size: 14px; margin-bottom: 0;'>برای قابلیت لاگ دائمی، نیاز به اتصال به دیتابیس آنلاین (مانند Firebase Firestore) است.</p>
        </div>
        """, unsafe_allow_html=True)


def render_chat_page():
    """Renders the main chat interface for regular users."""
    with st.sidebar:
        st.title(f"کاربر: {st.session_state.user_id}")
        
        st.markdown("---")
        if st.session_state.page_history:
            st.sidebar.button("🔙 بازگشت به صفحه قبلی", on_click=go_back, use_container_width=True)
        st.sidebar.button("🏠 بازگشت به صفحه اصلی", on_click=go_to_main_page, use_container_width=True)

        if st.sidebar.button("⚙️ حساب کاربری من", key="my_account_btn_chat", use_container_width=True):
            navigate_to("user_account")
        st.button("خروج از سیستم 🚪", on_click=logout, use_container_width=True)

    st.title("🧠 دستیار دانش هوشمند شرکت سپاهان")
    st.subheader("💡 سوالات خود را در مورد دستورالعمل‌ها و رویه‌های شرکت بپرسید.")

    st.info("💡 من اینجا هستم تا به سوالات شما بر اساس اسناد داخلی شرکت پاسخ دهم.")

    # Load knowledge base (and cache it)
    vector_store, _ = load_knowledge_base_from_index(aval_ai_api_key) # Use aval_ai_api_key

    if vector_store is None:
        st.error("🚨 پایگاه دانش بارگذاری نشد.")
        if st.session_state.is_admin:
            st.warning("به نظر می‌رسد پایگاه دانش خالی است یا با مشکل مواجه شده است. به عنوان مدیر، می‌توانید از بخش مدیریت، اسناد جدید را بارگذاری و پایگاه دانش را بازسازی کنید.")
            if st.button("🛠️ رفتن به پنل مدیریت برای بازسازی پایگاه دانش", use_container_width=True):
                navigate_to("admin")
        else:
            st.warning("لطفاً با مدیر سیستم تماس بگیرید تا اسناد را اضافه یا پایگاه دانش را بازسازی کند.")
        return # Stop execution if KB not loaded

    retriever = vector_store.as_retriever()
    
    # Initialize LLMs using CustomLLM
    llm = CustomLLM(api_key=aval_ai_api_key, model_name="your-text-model") # Use aval_ai_api_key
    multimodal_llm = CustomLLM(api_key=aval_ai_api_key, model_name="your-multimodal-model") # Use aval_ai_api_key

    # Setup RetrievalQA chain
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=retriever,
        return_source_documents=False # Set to True if you want to show sources
    )

    # --- User File/Image Upload for Context ---
    st.markdown("<h3 class='chat-section-header'>🖼️ افزودن فایل/تصویر به مکالمه</h3>", unsafe_allow_html=True)
    user_uploaded_context_file = st.file_uploader(
        "یک فایل PDF یا تصویر (JPG, PNG) برای افزودن به سوال فعلی آپلود کنید.",
        type=["pdf", "jpg", "jpeg", "png", "docx", "xlsx"],
        key="user_context_uploader",
    )

    # --- Chat Interface ---
    # Display chat messages from history on app rerun
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            if message["type"] == "text":
                st.markdown(message["content"])
            elif message["type"] == "image":
                st.image(message["content"], caption="تصویر آپلود شده", use_column_width=True)
                if "text_content" in message: # Display text description if available
                    st.markdown(message["text_content"])

    # Add a typing indicator placeholder
    if st.session_state.get("is_generating_response", False):
        with st.chat_message("assistant"):
            st.markdown('<div class="typing-indicator"><span>.</span><span>.</span><span>.</span></div>', unsafe_allow_html=True)


    # Accept user input
    if prompt := st.chat_input("سوال خود را در مورد دستورالعمل‌ها بپرسید..."):
        st.session_state.is_generating_response = True # Set flag when user inputs
        user_message_display = {"type": "text", "content": prompt}
        
        # Prepare prompt parts for LLM (text and/or image data)
        # For custom LLM, you might need to adjust how prompt_parts are structured based on its API
        llm_prompt_input_text = prompt # Default for text-only LLM
        llm_prompt_input_multimodal = None # For multimodal input

        # Handle uploaded file/image
        if user_uploaded_context_file:
            file_type = user_uploaded_context_file.type
            file_extension = os.path.splitext(user_uploaded_context_file.name)[1].lower()

            if file_extension == ".pdf":
                st.info("در حال پردازش PDF آپلود شده برای سوال فعلی...")
                temp_file_path = None
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                        temp_file.write(user_uploaded_context_file.getvalue())
                        temp_file_path = temp_file.name
                    
                    pdf_docs = process_file_to_documents(temp_file_path, file_extension)
                    pdf_text = "\n".join([doc.page_content for doc in pdf_docs])
                    
                    user_message_display["content"] += f"\n\n(متن مرتبط از فایل PDF: {pdf_text[:500]}...)"
                    llm_prompt_input_text += f"\n\nمتن مرتبط از فایل PDF: {pdf_text}" # Add PDF text to LLM input
                    st.success("PDF برای سوال فعلی پردازش شد.")
                except Exception as e:
                    st.error(f"خطا در پردازش PDF برای سوال فعلی: {e}")
                finally:
                    if temp_file_path and os.path.exists(temp_file_path):
                        os.remove(temp_file_path)

            elif file_extension == ".docx":
                st.info("در حال پردازش فایل Word آپلود شده برای سوال فعلی...")
                temp_file_path = None
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                        temp_file.write(user_uploaded_context_file.getvalue())
                        temp_file_path = temp_file.name
                    
                    doc_docs = process_file_to_documents(temp_file_path, file_extension)
                    doc_text = "\n".join([doc.page_content for doc in doc_docs])
                    
                    user_message_display["content"] += f"\n\n(متن مرتبط از فایل Word: {doc_text[:500]}...)"
                    llm_prompt_input_text += f"\n\nمتن مرتبط از فایل Word: {doc_text}" # Add DOCX text to LLM input
                    st.success("فایل Word برای سوال فعلی پردازش شد.")
                except Exception as e:
                    st.error(f"خطا در پردازش فایل Word برای سوال فعلی: {e}")
                finally:
                    if temp_file_path and os.path.exists(temp_file_path):
                        os.remove(temp_file_path)

            elif file_extension == ".xlsx":
                st.info("در حال پردازش فایل Excel آپلود شده برای سوال فعلی...")
                temp_file_path = None
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
                        temp_file.write(user_uploaded_context_file.getvalue())
                        temp_file_path = temp_file.name
                    
                    excel_docs = process_file_to_documents(temp_file_path, file_extension)
                    excel_text = "\n".join([doc.page_content for doc in excel_docs])
                    
                    user_message_display["content"] += f"\n\n(متن مرتبط از فایل Excel: {excel_text[:500]}...)"
                    llm_prompt_input_text += f"\n\nمتن مرتبط از فایل Excel: {excel_text}" # Add XLSX text to LLM input
                    st.success("فایل Excel برای سوال فعلی پردازش شد.")
                except Exception as e:
                    st.error(f"خطا در پردازش فایل Excel برای سوال فعلی: {e}")
                finally:
                    if temp_file_path and os.path.exists(temp_file_path):
                        os.remove(temp_file_path)

            elif "image" in file_type:
                st.info("در حال پردازش تصویر آپلود شده برای سوال فعلی...")
                base64_image = base64.b64encode(user_uploaded_context_file.getvalue()).decode('utf-8')
                user_message_display = {"type": "image", "content": f"data:{file_type};base64,{base64_image}", "text_content": prompt}
                
                # For multimodal LLM, prompt_parts might be a list of dicts like [{"text": "...", "inlineData": {"mimeType": "...", "data": "..."}}]
                llm_prompt_input_multimodal = [
                    {"text": prompt},
                    {"inlineData": {"mimeType": file_type, "data": base64_image}}
                ]
                st.success("تصویر برای سوال فعلی پردازش شد.")
            else:
                st.warning("فرمت فایل آپلود شده پشتیبانی نمی‌شود.")
        
        # Add user message to chat history and display
        st.session_state.messages.append({"role": "user", **user_message_display})
        with st.chat_message("user"):
            if user_message_display["type"] == "text":
                st.markdown(user_message_display["content"])
            elif user_message_display["type"] == "image":
                st.image(user_message_display["content"], caption="تصویر آپلود شده", use_column_width=True)
                st.markdown(user_message_display["text_content"])

        # Get assistant response
        with st.chat_message("assistant"):
            with st.spinner("🚀 دستیار هوش مصنوعی در حال پردازش سوال شماست..."):
                try:
                    full_response = ""
                    if user_uploaded_context_file and ("image" in file_type):
                        # For image input, use multimodal_llm directly with the structured input
                        # Ensure raw_response has a .content attribute, or adjust based on actual multimodal LLM output
                        raw_response = multimodal_llm._call(llm_prompt_input_multimodal) # Use _call for direct invocation
                        full_response = raw_response if isinstance(raw_response, str) else getattr(raw_response, 'content', "پاسخی دریافت نشد.")
                    else:
                        # For text, PDF, DOCX, XLSX input, use qa_chain (RAG)
                        response = qa_chain.invoke({"query": llm_prompt_input_text})
                        full_response = response["result"]

                    st.markdown(full_response)
                except Exception as e:
                    st.error(f"⚠️ متاسفانه در حال حاضر مشکلی در پردازش درخواست شما پیش آمده است. لطفاً لحظاتی دیگر دوباره تلاش کنید. (خطا: {e})")
                finally:
                    st.session_state.is_generating_response = False # Reset flag after generation

            assistant_message_content = {"role": "assistant", "type": "text", "content": full_response}
            st.session_state.messages.append(assistant_message_content)
        
        # JavaScript to scroll to the bottom after new messages are added
        st.markdown("""
            <script>
                var chatContainer = parent.document.querySelector('.main .block-container');
                if (chatContainer) {
                    chatContainer.scrollTop = chatContainer.scrollHeight;
                }
            </script>
        """, unsafe_allow_html=True)


    st.markdown("---")
    st.markdown(f"<p style='text-align: center; font-size: 13px; color: var(--subtle-text-color);'>نسخه آزمایشی v1.0 | تاریخ: {datetime.now().strftime('%Y-%m-%d')}</p>", unsafe_allow_html=True)
    st.markdown(f"<p style='text-align: center; font-size: 13px; color: var(--subtle-text-color);'>&copy; {datetime.now().year} گروه صنعتی سپاهان. تمامی حقوق محفوظ است.</p>", unsafe_allow_html=True)


def render_user_account_page():
    """Renders the user account management page (e.g., change password)."""
    with st.sidebar:
        st.title(f"حساب کاربری: {st.session_state.user_id}")
        st.markdown("---")
        if st.session_state.page_history:
            st.sidebar.button("🔙 بازگشت به صفحه قبلی", on_click=go_back, use_container_width=True)
        st.sidebar.button("🏠 بازگشت به صفحه اصلی", on_click=go_to_main_page, use_container_width=True)
        st.button("خروج از سیستم 🚪", on_click=logout, use_container_width=True)

    st.title("👤 مدیریت حساب کاربری")
    st.info(f"شما به عنوان **{st.session_state.user_id}** وارد شده‌اید.")

    st.subheader("تغییر رمز عبور")
    with st.form("change_password_form"):
        current_password = st.text_input("رمز عبور فعلی", type="password", key="current_pass_input")
        new_password = st.text_input("رمز عبور جدید", type="password", key="new_pass_input")
        confirm_new_password = st.text_input("تکرار رمز عبور جدید", type="password", key="confirm_new_pass_input")

        if st.form_submit_button("تغییر رمز عبور", use_container_width=True, type="primary"):
            users_data = load_users()
            user_type_key = "admin_users" if st.session_state.is_admin else "users"
            
            user_found = False
            for user_info in users_data.get(user_type_key, []):
                if user_info["username"] == st.session_state.user_id:
                    user_found = True
                    if user_info["password"] == current_password:
                        if new_password == confirm_new_password:
                            if new_password and len(new_password) >= 4: # Basic validation
                                user_info["password"] = new_password
                                save_users(users_data)
                                st.success("✅ رمز عبور با موفقیت تغییر یافت.")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.warning("⚠️ رمز عبور جدید حداقل باید 4 کاراکتر باشد.")
                        else:
                            st.warning("⚠️ رمز عبور جدید و تکرار آن مطابقت ندارند.")
                    else:
                        st.error("❌ رمز عبور فعلی اشتباه است.")
                    break
            if not user_found:
                st.error("خطا: کاربر جاری در پایگاه داده یافت نشد. لطفاً دوباره وارد شوید.")

    st.markdown("---")
    st.markdown(f"<p style='text-align: center; font-size: 13px; color: var(--subtle-text-color);'>&copy; {datetime.now().year} گروه صنعتی سپاهان. تمامی حقوق محفوظ است.</p>", unsafe_allow_html=True)


# --- Main App Flow Control ---
if st.session_state.authenticated:
    if st.session_state.current_page == "admin":
        render_admin_page()
    elif st.session_state.current_page == "chat":
        render_chat_page()
    elif st.session_state.current_page == "user_account":
        render_user_account_page()
    else:
        # Fallback for unexpected current_page values
        st.error("خطای ناوبری: صفحه نامعتبر.")
        logout() # Force logout to reset
else:
    render_login_page()
    st.markdown("<hr style='border-top: 1px solid var(--border-color); margin-top: 40px;'>", unsafe_allow_html=True) # Light gray line
    st.markdown(f"<p style='text-align: center; font-size: 13px; color: var(--subtle-text-color);'>&copy; {datetime.now().year} گروه صنعتی سپاهان. تمامی حقوق محفوظ است.</p>", unsafe_allow_html=True)
